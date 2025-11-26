import re

from docx import Document
from docx.shared import Pt


def add_markdown_paragraph(paragraph, text):
    """
    Process the given text for markdown bold markers (**...**) and add runs to the paragraph.
    """
    # Split the text on patterns that are enclosed with ** (capturing the delimiters)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        # If part is enclosed in **, add it as a bold run.
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def markdown_to_docx(document: Document, markdown_text: str):
    """
    Convert markdown-formatted string into a docx document.
    This converter processes:
      - Lines starting with "#", "##", "###" as headings.
      - Lines starting with "- " as bullet points.
      - Any text in a line, including markdown bold (**text**), into a paragraph.
    """
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- "):
            p = document.add_paragraph(style='List Bullet')
            add_markdown_paragraph(p, stripped[2:].strip())
        else:
            # For normal text, add a paragraph and process for bold markers.
            p = document.add_paragraph()
            add_markdown_paragraph(p, stripped)

def write_docx(filename: str, content: str, use_markdown: bool = False):
    """
    Save content to a DOCX file. If use_markdown is True, try to adapt markdown formatting.
    """
    doc = Document()
    # Optional: set default font size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    if use_markdown:
        markdown_to_docx(doc, content)
    else:
        doc.add_paragraph(content)
    doc.save(filename)